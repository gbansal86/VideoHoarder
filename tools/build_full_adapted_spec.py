from copy import deepcopy
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path(r"C:\Users\gaurav.bansal\Downloads\Codex_Master_Spec_ChatGPT_Video_Intelligence_CHATGPT_ONLY.docx")
OUT = Path(r"C:\Users\gaurav.bansal\Documents\ChatGPT\VideoHoarder\VideoHoarder_FULL_ChatGPT_Video_Intelligence_Master_Spec.docx")

NAVY="17324D"; BLUE="246BCE"; TEAL="178C8C"; PALE="E8F6F5"; AMBER="B54708"; RED="B42318"; GRAY="667085"

def set_font(run, name="Aptos", size=9.2, bold=False, color="101828"):
    run.font.name=name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"),name)
    run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=RGBColor.from_string(color)

def shade_paragraph(p, fill=PALE):
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),fill); pPr.append(shd)
    pPr.append(OxmlElement("w:keepNext"))
    p.paragraph_format.left_indent=Inches(.12); p.paragraph_format.right_indent=Inches(.12)
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(5)

def insert_after(paragraph, label, status, text):
    new_p=OxmlElement("w:p"); paragraph._p.addnext(new_p)
    p=paragraph._parent.add_paragraph(); p._p.getparent().remove(p._p); new_p.addnext(p._p); new_p.getparent().remove(new_p)
    shade_paragraph(p)
    r=p.add_run(f"VIDEOHOARDER ADAPTATION — {status.upper()}  "); set_font(r,bold=True,color=TEAL,size=8.8)
    r=p.add_run(text); set_font(r,size=8.8)
    return p

notes = {
1:("Modified","Keep the manual ChatGPT package/import architecture, but preserve VideoHoarder’s existing local Ollama Ask Local AI and embedding support as a separate local-only feature. ChatGPT remains the authoritative engine only for the reviewed package/import intelligence workflow; no direct OpenAI API upload is introduced."),
2:("Required first step","Audit app/app.py before implementation. Reuse Phase 2 package lifecycle, Phase 3 validation, Phase 4 cards, Phase 5 search/reports, Phase 6 retrieval/collections, comment intelligence, artifact registry, safe rename, move preview/undo and Clip Studio functions. Do not build parallel replacements."),
3:("New normalized layer","Add a versioned v3 normalized model above existing summary and taxonomy schemas. Legacy detailed-summary, taxonomy-only and JSONL results continue through adapters; projections feed existing canonical summaries, content cards, reports, search documents and chunks."),
4:("Mostly existing; extend","VIDEO_ID matching, original title preservation and meaningful English titles already exist. Add short_title, explicit source_language/title_translated, field provenance and manual locks. Physical renaming already exists through apply_chatgpt_titles_and_rename() and must be reused after approval."),
5:("Partially existing","Current taxonomy packages and Phase 5 canonical topic/tag processing cover the base. Unify taxonomy export, add stable synonyms/questions/not_about handling, and show before/after taxonomy differences before database updates or organization moves."),
6:("New consolidation","Phase 5 already normalizes tags and derives categories, but no complete canonical entity registry exists. Add versioned taxonomy dictionaries with aliases and multilingual forms; migration must preserve existing accepted terms rather than recategorizing the library blindly."),
7:("Partially existing","Phase 4 detects content types. Expand the controlled vocabulary and allow multiple types while maintaining compatibility with existing content-card and report rendering."),
8:("Existing foundation; extend","Detailed summaries and timestamped sections already exist. Preserve validate_and_normalize_sections(), Phase 3 timestamp analysis and report timestamp links; add executive/detailed summaries, topic-change chapters, timestamp coverage and multipart handling without silent truncation."),
9:("New extraction rule","Add named-item completeness fields to the unified prompt/schema and content-card projection. Existing entity/ingredient/herb/product fields are reused, but stated counts must be reconciled without inventing missing items."),
10:("New enrichment","Phase 5/6 retrieval can find multiple chunks; formalize topic_occurrences in normalized intelligence and update Find Timestamp to return all meaningful ranges from the same VIDEO_ID."),
11:("New report index","Generate Things Discussed from normalized entities, topics, cards, claims and warnings. Add it to the per-video report with local clickable timestamps and hide empty groups."),
12:("New structured fields","Extend normalized intelligence and reports with concrete takeaways, attributed facts and exact numbers. Preserve source wording, units, timestamps and speaker attribution."),
13:("Partially existing","Existing summaries extract several entity lists. Expand to the specified typed entity object and canonical aliases; never promote inferred or comment-only entities into transcript facts."),
14:("Existing foundation; extend","Phase 4 supports how-to content and existing summary sections carry steps. Normalize requirements, prerequisites, ordered steps, warnings and expected result with evidence and timestamps."),
15:("Existing foundation; safety upgrade","Remedy cards already exist in the detailed-summary schema. Add claim attribution, verification state, exact stated dosage/timing and evidence; never manufacture generic medical warnings as source statements."),
16:("Existing foundation; extend","Recipe cards already exist. Expand normalization for servings, equipment, substitutions, storage, time and temperature while retaining only source-supported values."),
17:("Partially existing","Existing prompts request comparisons/rankings and Phase 4 detects structured content. Normalize criteria, ordered/unordered distinction, attribution, evidence and timestamps."),
18:("Partially existing","Existing product/entity extraction is reused. Add typed product/service/tool fields and preserve price/recommendation attribution and timestamp."),
19:("New evidence discipline","Existing summaries contain claims and safety notes but need stronger provenance. Add verification states, speaker/source attribution and explicit source-only reference capture."),
20:("Partially existing","Existing schema supports glossary and Q&A. Add evidence, timestamp and confidence, then project them into reports and search without answering beyond supplied evidence."),
21:("Existing foundation; extend","Comment downloads, meaningfulness scoring, meaningful-comments transcript and report injection already exist. Preserve build_comment_intelligence(), keep comments separate, add analyzed counts/themes/cautions and provenance IDs."),
22:("New mandatory layer","Add stable evidence IDs during package construction and validate returned evidence references against exported evidence. Keep transcript, metadata, existing intelligence and viewer comments as separate source types."),
23:("Prompt and validator change","Apply these rules to all four package modes and reject prohibited destructive instructions. Empty/uncertain output is preferred to guessed timestamps, quantities, names, sources or steps."),
24:("Partially existing","Transcript availability is already tracked. Expand statuses to AVAILABLE/PARTIAL/MISSING/UNUSABLE, record classification basis and quality coverage, and prevent metadata-only results from claiming transcript-derived detail."),
25:("Modified for privacy","Use a shared evidence builder over existing payload functions. Add a visible preflight listing every included field, sizes and sensitive-content warnings; exclude media, thumbnails, local paths, database, logs, cookies, credentials and unrelated files."),
26:("Modified to four modes","Implement MASTER, FOCUSED, TAXONOMY and SAME_TOPIC under one envelope. Map existing detailed summaries, taxonomy/tag cleanup and full/focused collection packages into the shared lifecycle."),
27:("New package manager capability","Replace arbitrary truncation for master analysis with size estimates, per-video parts and a merge package. Taxonomy/focused modes may use explicitly disclosed excerpts; manifests record omissions and source hashes."),
28:("New version registry","Extend Phase 2 manifests/history with package, schema and prompt versions, source hashes, part relationships and required result filename. Use collision-resistant package IDs rather than only second-level timestamps."),
29:("Existing foundation; major extension","Reuse Phase 2 import, manifest comparison, retries and archives plus Phase 3 checks. Add raw immutable archive, normalized result, evidence validation, per-field diff, locks and per-VIDEO_ID atomic transactions."),
30:("New review layer","Folder move preview/undo already exists, but general field-level preview does not. Add current/proposed value, evidence source, confidence, lock state and selected apply action before title/taxonomy/intelligence changes."),
31:("Partially existing","VIDEO_ID and canonical filenames provide stability. Add normalized-result hashes and transaction IDs so repeated imports do not duplicate cards, sections or history."),
32:("Modified","Retain data/chatgpt packages/results/retry/archive and per-video _data canonical summaries. Add raw, normalized, rejected, previews and transactions subfolders plus provenance/intelligence sidecars; migrate by copying, never deleting originals."),
33:("Existing foundation; extend","Reuse Phase 4 content detection/extraction and CSV exports. Project normalized guides, remedies, recipes, comparisons, rankings, products, claims and warnings into versioned cards with evidence."),
34:("Existing foundation; extend","Reuse Phase 5 dashboards/topic pages/search and Phase 6 collections/retrieval. Refresh only changed VIDEO_IDs and expose provenance/quality without creating another independent Knowledge Center."),
35:("Modified","For ChatGPT, generate a FOCUSED manual-upload package from selected videos; the app does not call ChatGPT automatically. Keep existing Ollama evidence Q&A as Ask Local AI, clearly separate from imported authoritative intelligence."),
36:("Existing foundation; extend","Reuse phase5_find_timestamp() and Phase 6 chunk search/semantic retrieval. Return multiple occurrences, evidence reason, local seek links and confidence; never fabricate timestamps."),
37:("Existing foundation; extend","Reuse phase5_related_for_video(), web_local_related() and semantic signals. Add duplicate collapse, novelty/diversity, unique contribution and user include/exclude controls."),
38:("Partially existing; advanced addition","Existing duplicate audits/topic groups are starting points. Add stable duplicate groups, side-by-side review and a canonical recommendation. General Mark Delete/delete controls exist, but integrated duplicate deletion remains manual and confirmed."),
39:("Partially existing","Collections and full/focused packages already provide cross-video inputs. Add separate normalized synthesis storage for agreements, contradictions and unique contributions; never overwrite per-video master intelligence."),
40:("Existing foundation; extend","Reuse Phase 5 topic pages and collection summaries. Add canonical entity dossiers/relationships only after the registry and provenance model are stable; avoid an expensive full graph rebuild on page load."),
41:("Partially existing","English translation/normalization is already requested and language checks exist. Add source-language aliases, translated-title metadata and multilingual search fields without overwriting original text."),
42:("Existing foundation; extend","Reports already show ChatGPT titles, summaries, timestamped sections, local video links and meaningful comments. Add unified sections, Things Discussed, evidence labels, quality state and optional clearly separated ChatGPT perspective."),
43:("Existing foundation; extend","The desktop/API already exposes rich video details. Add package/import status, normalized schema version, locks, provenance, stale state, duplicate group and related/synthesis links."),
44:("Existing foundation; make selective","Report regeneration, Phase 4 extraction and Phase 5/6 rebuild functions exist. Orchestrate them after commit using source hashes so unchanged artifacts are not rebuilt."),
45:("New integrity layer","Use transcript, description, comments, summary and prompt/schema hashes to mark only affected intelligence/report/search/chunk artifacts stale. Surface stale reasons and safe refresh actions."),
46:("New lineage metadata","Record package ID, result hash, prompt/schema version, evidence hashes, import transaction, normalized hash and downstream artifact versions per VIDEO_ID."),
47:("New UI","Build review queues from existing Phase 2/3 status and validation reports. Show missing/failed/partial/stale/low-confidence videos and actionable retry/preview controls."),
48:("Deferred opt-in workflow","Do not mix web facts into source-only extraction. Any future verification package must be separately requested, labeled, stored and reviewed."),
49:("Modified existing page","Keep the native ChatGPT Processing page and maintenance buttons. Reorganize into Select, Review Evidence, Create Package, Manual Exchange, Import, Apply and Maintenance panels with truthful statuses."),
50:("Required UI documentation","Add concise hover help explaining local processing versus manual upload, package modes, privacy, evidence, preview/apply/undo, rename, folder movement, duplicates and Clip Studio handoff."),
51:("Existing foundation; harden","Reuse error logging, audits and package status. Add precise states, per-video rollback, validation reports, recoverable retry and no false “sent” or “completed” status."),
52:("Expanded tests","Preserve existing tests and feature audit. Add permanent fixtures for old/new schemas, multilingual, missing/partial transcript, structured content, comments, repeated topics, Unicode, multipart, duplicate and transaction failure."),
53:("Existing concern; extend","Preserve PyInstaller/runtime packaging and path-length protections. Ensure schemas/prompts/migrations/fixtures are bundled and no source DOCX, secrets, local data or QA artifacts enter the release."),
54:("Updated deliverables","Deliver code, numbered migrations, schemas/prompts, regression fixtures, UI/help changes, migration/audit report and a verified Windows build. This document defines target scope; implementation follows phases."),
55:("Backlog with prioritization","Retain valuable additions but gate them behind the core schema/export/import/projection work. Features that depend on provenance, duplicates or atomic imports are not implemented prematurely."),
56:("Modified priority","Priority 1: compatibility and normalized schema. 2: shared export/preflight. 3: raw archive/validation/preview/atomic import. 4: downstream projections and rename/move integration. 5: same-topic, duplicates and Clip Studio plan handoff. 6: scaling/portability."),
57:("Final architectural outcome","VideoHoarder remains a desktop library manager using transparent local files for manual ChatGPT processing. Existing downloads, reports, maintenance, search, collections, rename/move and Ollama features remain intact."),
58:("New quality extensions","Implement subsections through the normalized model and evidence IDs. Source priority, novelty, corrections, time sensitivity, technical extraction and structured tables must not bypass source-only rules."),
59:("New enrichment","Add playlist/series/episode fields from official metadata and source evidence. Preserve ordering and VIDEO_ID identity; do not infer missing episode relationships solely from similar titles."),
60:("Mostly UI reuse","Existing thumbnails are already used in library/related views. Do not send thumbnails for ChatGPT image analysis in this scope; use local thumbnail references only in the UI."),
61:("New database/UI capability","Add per-field source, confidence, updated_at and manual lock. Import preview must expose attempted locked changes; explicit override is required and previous values remain auditable."),
62:("Extend existing audits","Reuse artifact registry, Phase 0/1 audits and path repair utilities. Add read-only media/DB/transcript/result/report/intelligence/collection/hash orphan checks; repairs remain separate and never auto-delete."),
63:("New reviewed recommendation","Score canonical duplicate candidates using completeness, media/transcript quality, metadata/comments and source authenticity. The user chooses; selection is reversible and does not delete variants."),
64:("New ranking policy","Enhance existing related-video outputs to collapse duplicate groups and reward complementary subtopics, warnings, approaches and examples. Provide include-variants control."),
65:("New reliability layer","Wrap each VIDEO_ID import in a transaction covering DB and staged file changes. Record STARTED/COMMITTED/ROLLED_BACK/FAILED, retain raw JSON and allow per-video undo without media deletion."),
66:("New permanent fixtures","Build fixtures from sanitized representative old/new shapes and run them for every schema/importer change. Include exact expected normalized output, not only parse success."),
67:("New hardening","Add crash-safe per-VIDEO_ID leases for download, import, report rebuild, rename/move, deletion and migration. Read-only viewing remains available; stale locks recover safely."),
68:("New foundation before DB expansion","Replace scattered future ALTER TABLE changes with numbered migrations, backup, validation and history. Existing databases upgrade idempotently and remain recoverable."),
69:("Hardening phase","Use existing incremental hashes/indexes and avoid all-pairs full-transcript work during UI requests. Add queue progress/cancel/retry and telemetry; test 100, 1,000 and 10,000+ records."),
70:("Extend existing portability","Build on portable metadata backup and artifact manifests. Export a per-VIDEO_ID record with schema versions/hashes and no cookies, keys, private browser data or fragile absolute paths."),
71:("Already a governing rule; enforce","Existing move preview/undo and manual deletion controls align with this. Validators must ignore/reject AI deletion authority; title rename, folder move, clip merge, archive and delete require explicit user action."),
72:("New advanced workflow on existing components","Use collections/search/topic pages plus Phase 6 retrieval to create SAME_TOPIC packages. Validate returned VIDEO_ID/ranges and store synthesis separately. Add Send Merge Plan to Clip Studio on top of existing VIDEO_ID-based web_merge_clips(), never automatic cutting."),
73:("New optimization over existing comments","Use transcript/metadata/intelligence to recommend whether comments are useful before download/package. Keep manual override and never discard already downloaded comments because of a later low-usefulness score."),
74:("Required expanded tests","Add these cases to automated fixtures and UI/runtime acceptance tests, including existing rename/move behavior, manual duplicate decisions and validated Clip Studio handoff."),
75:("Modified final flow","Per-video and cross-video flows remain manual ChatGPT exchanges. Rename and taxonomy folder movement reuse existing reviewed functions; duplicate deletion and clip execution remain user-authorized operations."),
76:("Optional; new separated field","Allow PS only inside structured JSON and label it NOT SOURCE-EXTRACTED FACT. It never overwrites evidence-backed fields or outranks transcript evidence in search."),
77:("Coverage statement","This full adapted document preserves the complete baseline while resolving conflicts with the real VideoHoarder architecture. The shorter adapted brief remains an executive summary; this is the implementation authority."),
}

doc=Document(SRC)

# Update cover without destroying source structure.
if doc.paragraphs:
    p=doc.paragraphs[0]
    p.text="VideoHoarder\nFull ChatGPT Video Intelligence Master Specification"
    for i,r in enumerate(p.runs): set_font(r,name="Aptos Display",size=24 if i==0 else 22,bold=True,color=NAVY)
if len(doc.paragraphs)>1:
    doc.paragraphs[1].text="Complete 77-Section Baseline + VideoHoarder Extensions • Existing / Modified / New / Deferred Decisions"
    for r in doc.paragraphs[1].runs:set_font(r,size=11,bold=True,color=BLUE)
if len(doc.paragraphs)>2:
    doc.paragraphs[2].text=("Purpose: This is the full implementation authority for VideoHoarder. It preserves the complete original ChatGPT video-intelligence baseline and adds an explicit VideoHoarder adaptation beneath every numbered section. Manual file exchange remains the ChatGPT workflow; existing application functionality is preserved and reused.")

# Insert front-matter status legend before Section 1.
first_h1=next((p for p in doc.paragraphs if p.style.name=="Heading 1"),None)
if first_h1:
    intro=OxmlElement("w:p"); first_h1._p.addprevious(intro)
    p=first_h1._parent.add_paragraph(); p._p.getparent().remove(p._p); intro.addnext(p._p); intro.getparent().remove(intro)
    shade_paragraph(p,"EAF2FC")
    r=p.add_run("HOW TO READ THIS DOCUMENT  ");set_font(r,bold=True,color=BLUE,size=9)
    r=p.add_run("Each numbered source section is preserved. The teal VideoHoarder block immediately below its heading records whether the requirement already exists, must be extended or modified, is new, or is deferred. When existing code is named, it should be reused rather than duplicated.");set_font(r,size=9)

# Add adaptation immediately after every top-level numbered heading.
for p in list(doc.paragraphs):
    if p.style.name != "Heading 1": continue
    m=re.match(r"^(\d+)\.",p.text.strip())
    if not m: continue
    n=int(m.group(1)); status,text=notes[n]
    insert_after(p, f"Section {n}", status, text)

# VideoHoarder-specific extension requested after the original 77-section baseline.
doc.add_heading("78. ChatGPT Request History and Per-Video Feature Coverage Registry", level=1)
p=doc.add_paragraph()
shade_paragraph(p)
r=p.add_run("VIDEOHOARDER EXTENSION — NEW CORE CAPABILITY  ");set_font(r,bold=True,color=TEAL,size=8.8)
r=p.add_run("Track exactly what was sent to ChatGPT, what was requested, what was returned, and which intelligence features are complete, missing, stale or failed for every VIDEO_ID. This enables a new feature to be requested only for pre-existing videos that do not already have that feature.");set_font(r,size=8.8)

doc.add_heading("78.1 What Must Be Recorded for Every Outgoing Package", level=2)
for text in [
    "package_id, package revision, package mode, schema version, prompt version, creation time and user-selected scope.",
    "Exact VIDEO_ID list plus a per-video source snapshot hash.",
    "Exact requested feature IDs, feature schema versions and required output fields.",
    "Which evidence types were included: metadata, description, transcript, timestamped sections, chapters, comments and existing intelligence.",
    "Character/token estimates, transcript/comment coverage, multipart/continuation information and explicitly omitted evidence.",
    "SHA-256 hash of every outgoing package file and the readable PACKAGE_README instructions.",
    "Manual exchange state: CREATED, READY_FOR_UPLOAD, AWAITING_RESULT, RESULT_DETECTED, VALIDATED, APPLIED, PARTIAL, FAILED, RETRY_REQUIRED or ARCHIVED.",
]: doc.add_paragraph(text,style="List Bullet")

doc.add_heading("78.2 What Must Be Recorded for Every Returned Result", level=2)
for text in [
    "Raw result filename/hash and immutable archive location.",
    "Returned package_id, VIDEO_IDs, schema/prompt version and requested feature IDs.",
    "Per-feature outcome for each VIDEO_ID: COMPLETE, PARTIAL, EMPTY_NOT_APPLICABLE, MISSING, INVALID, FAILED or STALE.",
    "Normalized-result hash, validation warnings/errors, evidence coverage and import transaction ID.",
    "Fields accepted, rejected, skipped because locked, or left unchanged.",
    "Downstream artifacts refreshed and their resulting hashes/versions.",
]: doc.add_paragraph(text,style="List Bullet")

doc.add_heading("78.3 Stable Feature IDs and Coverage States", level=2)
doc.add_paragraph("Every independently requestable capability receives a stable feature ID and version. Examples:")
code=doc.add_paragraph()
code.style=doc.styles["Normal"]
code.paragraph_format.left_indent=Inches(.25)
run=code.add_run("""title_en.v1
taxonomy.v3
executive_summary.v2
detailed_summary.v3
timeline_sections.v3
topic_occurrences.v1
things_discussed.v1
entities.v2
how_to_guides.v2
remedies.v2
recipes.v2
claims_warnings.v2
viewer_comment_intelligence.v2
evidence_provenance.v1
same_topic_clip_plan.v1""")
set_font(run,name="Consolas",size=8.5,color="344054")
doc.add_paragraph("Coverage is stored by VIDEO_ID + feature_id + feature_version + source_snapshot_hash. A feature is reusable only when its version is compatible and the relevant source snapshot has not changed.")

doc.add_heading("78.4 Incremental New-Feature Workflow for Existing Videos", level=2)
steps=[
    "Register the new feature ID, schema version, required evidence types, output schema and merge policy.",
    "Query the Feature Coverage Registry for videos where the feature is MISSING, PARTIAL, FAILED, incompatible, or STALE.",
    "Show the candidate count and allow filtering by channel, category, collection, date, transcript status or explicit selection.",
    "Build a FEATURE_ENRICHMENT package containing only the new feature request and the minimum evidence required for it.",
    "Do not ask ChatGPT to regenerate already accepted title, taxonomy, summary or other unrelated fields.",
    "Validate that the returned result contains only authorized feature fields and exact VIDEO_IDs.",
    "Merge the new feature into existing normalized intelligence without overwriting unrelated accepted fields.",
    "Mark per-video coverage and selectively refresh only downstream artifacts that consume the new feature.",
    "Create retry packages only for videos/features that remain missing or invalid.",
]
for i,s in enumerate(steps,1):
    p=doc.add_paragraph();r=p.add_run(f"Step {i} — ");set_font(r,bold=True,color=BLUE,size=9.2);p.add_run(s)

doc.add_heading("78.5 Required Storage Model", level=2)
code=doc.add_paragraph()
run=code.add_run("""chatgpt_requests
  request_id, package_id, mode, schema_version, prompt_version,
  requested_features_json, package_hash, created_at, status

chatgpt_request_videos
  request_id, video_id, source_snapshot_hash,
  evidence_manifest_json, requested_features_json

video_feature_coverage
  video_id, feature_id, feature_version, status,
  source_snapshot_hash, package_id, result_hash,
  evidence_coverage, completed_at, stale_reason, last_error

chatgpt_import_field_history
  transaction_id, video_id, feature_id, field_path,
  previous_hash, proposed_hash, action, lock_state, timestamp""")
set_font(run,name="Consolas",size=8.3,color="344054")

doc.add_heading("78.6 Source-Aware Staleness", level=2)
for text in [
    "A comments-only change must not automatically stale transcript-only features.",
    "A transcript replacement stales transcript-derived summaries, timeline, cards, claims and evidence, but not authoritative metadata fields that did not change.",
    "A prompt improvement does not force all videos to reprocess unless its feature version is declared incompatible or the user chooses an upgrade campaign.",
    "Manual edits remain authoritative and locked fields are never made stale merely because ChatGPT has a newer proposal.",
]: doc.add_paragraph(text,style="List Bullet")

doc.add_heading("78.7 UI Requirements", level=2)
for text in [
    "Package History: show what was sent, what was requested, included evidence, file hashes, result state and imported fields.",
    "Feature Coverage Matrix: rows are videos; columns are stable feature IDs; cells show Complete, Missing, Partial, Failed, Stale or Not Applicable.",
    "Create Package for Missing Feature: select a feature and automatically target only eligible pre-existing videos.",
    "Upgrade Feature Version: preview how many videos are compatible, stale or candidates for optional reprocessing.",
    "Per-video ChatGPT History: display every package/result/feature change without exposing raw local paths or secrets in outgoing files.",
]: doc.add_paragraph(text,style="List Bullet")

doc.add_heading("78.8 Example", level=2)
doc.add_paragraph("If remedies.v2 is introduced after 5,000 videos have already been summarized, VideoHoarder queries only videos whose content type or transcript evidence makes remedies relevant and whose remedies.v2 coverage is not complete. The package asks only for remedies.v2 plus evidence provenance. Returned remedies are merged into each video’s intelligence; accepted titles, taxonomy, summaries and other cards remain unchanged.")

doc.add_heading("78.9 Acceptance Tests", level=2)
for i,text in enumerate([
    "A package history record reproduces the exact requested features and evidence manifest for an earlier manual upload.",
    "Adding a new feature selects only videos missing that compatible feature version.",
    "A feature-only result cannot overwrite title, taxonomy or summary fields that were not requested.",
    "A transcript change marks only dependent features stale.",
    "Partial results generate retries only for missing VIDEO_ID + feature pairs.",
    "Repeated import of the same result is idempotent.",
    "Feature coverage survives application restart, migration, folder rename and category movement because identity is VIDEO_ID-based.",
],1):
    p=doc.add_paragraph();r=p.add_run(f"Test {i} — ");set_font(r,bold=True,color=BLUE,size=9.2);p.add_run(text)

doc.add_heading("79. Final ChatGPT Processing UI Outlook", level=1)
p=doc.add_paragraph()
shade_paragraph(p)
r=p.add_run("VIDEOHOARDER DESIGN REFERENCE — APPROVED FUNCTIONAL PLACEMENT  ");set_font(r,bold=True,color=TEAL,size=8.8)
r=p.add_run("This mockup defines the intended information architecture for ChatGPT Processing. It is a functional layout reference, not a requirement to replace the existing VideoHoarder visual design system.");set_font(r,size=8.8)

mockup = r"C:\Users\gaurav.bansal\Documents\ChatGPT\VideoHoarder\tools\chatgpt-processing-final-outlook-cropped.png"
pic_p=doc.add_paragraph();pic_p.alignment=WD_ALIGN_PARAGRAPH.CENTER
pic_p.add_run().add_picture(mockup,width=Inches(6.65))
cap=doc.add_paragraph("Figure 1 — Proposed final ChatGPT Processing workspace overview")
cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
for run in cap.runs:set_font(run,size=8,color=GRAY)

doc.add_heading("79.1 Navigation and Functional Placement", level=2)
placements=[
    ("Overview", "Workflow status, truthful manual-exchange notice, current packages, missing-feature campaigns and quick actions."),
    ("Create Package", "MASTER, FOCUSED, TAXONOMY, SAME_TOPIC and FEATURE_ENRICHMENT modes; selection, evidence scope, privacy preflight and estimates."),
    ("Request History", "Exact package/prompt/schema/features/evidence sent, result file, validation, import and archive history."),
    ("Feature Coverage", "Per-VIDEO_ID feature matrix and Create Package for Missing Feature workflow."),
    ("Import & Validate", "Inbox scan, immutable raw archive, schema/VIDEO_ID/evidence/timestamp checks and retry generation."),
    ("Review & Apply", "Field-level diff, locks, title approval, existing physical rename, report refresh and separate organization-move preview."),
    ("Duplicates", "Duplicate groups, recommended/manual canonical selection and explicit Keep/Archive/Mark Delete/Delete actions."),
    ("Clip Plans", "Validated ChatGPT merge plan, clip inclusion/order preview and manual Send to Clip Studio."),
    ("Maintenance", "Legacy folder sync, Phase 2 integrity/completeness audit and local comment-intelligence rebuild."),
]
for title,text in placements:
    p=doc.add_paragraph();r=p.add_run(title+" — ");set_font(r,bold=True,color=BLUE,size=9.2);p.add_run(text)

doc.add_heading("79.2 Interaction Rules", level=2)
for text in [
    "The default Overview must clearly state that VideoHoarder does not automatically upload data to ChatGPT.",
    "A user can move through Select → Review Evidence → Create Package → Manual Exchange → Import & Validate → Review & Apply without losing package context.",
    "Physical rename, folder movement, Clip Studio execution, archive and deletion remain distinct reviewed actions.",
    "Package History and Feature Coverage use VIDEO_ID and stable feature IDs so folder or filename changes do not break tracking.",
    "Maintenance tools remain accessible but visually separated from the normal package workflow.",
]: doc.add_paragraph(text,style="List Bullet")

doc.add_page_break()
doc.add_heading("80. Selected ChatGPT Processing UI Design", level=1)
p=doc.add_paragraph()
shade_paragraph(p)
r=p.add_run("SELECTED DESIGN - SIDEBAR COMMAND CENTER  ");set_font(r,bold=True,color=TEAL,size=8.8)
r=p.add_run("This is the approved primary application shell for ChatGPT Processing. The four non-selected concepts are retained in a separate design-alternatives document for future reference.");set_font(r,size=8.8)

concepts=[
    (1,"Sidebar Command Center","Best all-purpose desktop layout and closest to the existing proposed outlook.","Persistent navigation makes every function easy to find; overview and advanced tools coexist clearly.","The left navigation consumes horizontal space and the overview can become busy as more features are added."),
]
for idx,title,best,strength,tradeoff in concepts:
    if idx>1:doc.add_page_break()
    doc.add_heading(f"80.{idx} Design {idx} — {title}",level=2)
    img=rf"C:\Users\gaurav.bansal\Documents\ChatGPT\VideoHoarder\tools\chatgpt-processing-concept-{idx}.png"
    pp=doc.add_paragraph();pp.alignment=WD_ALIGN_PARAGRAPH.CENTER;pp.add_run().add_picture(img,width=Inches(6.65))
    cp=doc.add_paragraph(f"Figure {idx+1} — {title}");cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for rr in cp.runs:set_font(rr,size=8,color=GRAY)
    for label,value in [("Best for",best),("Strength",strength),("Tradeoff",tradeoff)]:
        p=doc.add_paragraph();r=p.add_run(label+" — ");set_font(r,bold=True,color=BLUE,size=9.2);p.add_run(value)

doc.add_heading("80.2 Implementation Direction",level=2)
doc.add_paragraph("Implement the Sidebar Command Center as the primary ChatGPT Processing shell. Keep its persistent navigation, manual-exchange notice, six-stage workflow, status counters, current-work list and quick actions. The four alternate concepts are intentionally excluded from this master specification and preserved separately as optional future references.")

# Styling and page furniture.
sec=doc.sections[0]
sec.top_margin=Inches(.76);sec.bottom_margin=Inches(.76);sec.left_margin=Inches(.86);sec.right_margin=Inches(.86)
header=sec.header.paragraphs[0];header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
header.text="VIDEOHOARDER | FULL CHATGPT VIDEO INTELLIGENCE MASTER SPEC"
for r in header.runs:set_font(r,size=7.7,bold=True,color=GRAY)
footer=sec.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
footer.text="Full VideoHoarder adaptation with Sections 78–80 • Manual ChatGPT package exchange • August 2026"
for r in footer.runs:set_font(r,size=7.5,color=GRAY)

for name,size,color in [("Heading 1",14,NAVY),("Heading 2",11.5,BLUE)]:
    st=doc.styles[name];st.font.name="Aptos Display";st.font.size=Pt(size);st.font.bold=True;st.font.color.rgb=RGBColor.from_string(color)
    st.paragraph_format.keep_with_next=True;st.paragraph_format.space_before=Pt(10);st.paragraph_format.space_after=Pt(4)
normal=doc.styles["Normal"];normal.font.name="Aptos";normal.font.size=Pt(10);normal.paragraph_format.space_after=Pt(4);normal.paragraph_format.line_spacing=1.06

doc.core_properties.title="VideoHoarder Full ChatGPT Video Intelligence Master Specification"
doc.core_properties.subject="Complete 77-section app-specific adaptation"
doc.core_properties.author="VideoHoarder Project"
doc.save(OUT)
print(OUT)
