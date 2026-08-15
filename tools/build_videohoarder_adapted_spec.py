from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\gaurav.bansal\Documents\ChatGPT\VideoHoarder\VideoHoarder_ChatGPT_Video_Intelligence_Adapted_Spec.docx")

NAVY = "17324D"
BLUE = "246BCE"
TEAL = "178C8C"
LIGHT_BLUE = "EAF2FC"
LIGHT_TEAL = "E8F6F5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
RED = "B42318"
AMBER = "B54708"
GREEN = "067647"
WHITE = "FFFFFF"
BLACK = "101828"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_font(run, name="Aptos", size=10.5, bold=False, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_keep(paragraph, keep_next=False, keep_lines=True):
    pf = paragraph.paragraph_format
    pf.keep_with_next = keep_next
    pf.keep_together = keep_lines


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, label, text, color=BLUE, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.45)
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_font(r, size=10, bold=True, color=color)
    r = p.add_run(text)
    set_font(r, size=10, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        shade(cell, NAVY)
        set_cell_margins(cell, 100, 110, 100, 110)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, size=font_size, bold=True, color=WHITE)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2:
                shade(cells[i], "F8FAFC")
            set_cell_margins(cells[i], 90, 110, 90, 110)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            for j, part in enumerate(str(value).split("\n")):
                if j:
                    p.add_run("\n")
                r = p.add_run(part)
                set_font(r, size=font_size, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.45)
    cell = table.cell(0, 0)
    shade(cell, "F7F8FA")
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, name="Consolas", size=8, color="344054")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.28)
section.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.08

for name, size, color, before, after in [
    ("Heading 1", 16, NAVY, 15, 6),
    ("Heading 2", 12.5, BLUE, 11, 4),
    ("Heading 3", 10.8, TEAL, 8, 3),
]:
    st = styles[name]
    st.font.name = "Aptos Display"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for lname in ["List Bullet", "List Bullet 2", "List Number"]:
    st = styles[lname]
    st.font.name = "Aptos"
    st.font.size = Pt(10.2)
    st.paragraph_format.space_after = Pt(2.5)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = hp.add_run("VIDEOHOARDER  |  CHATGPT VIDEO INTELLIGENCE")
set_font(r, size=8.5, bold=True, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("Adapted implementation specification • Manual ChatGPT package exchange • August 2026")
set_font(r, size=8, color=MID_GRAY)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(48)
p.paragraph_format.space_after = Pt(8)
r = p.add_run("VideoHoarder")
set_font(r, name="Aptos Display", size=31, bold=True, color=NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("ChatGPT Video Intelligence\nAdapted Master Implementation Specification")
set_font(r, name="Aptos Display", size=20, bold=True, color=BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(24)
r = p.add_run("A VideoHoarder-specific revision of the ChatGPT-only master specification")
set_font(r, size=11.5, color=MID_GRAY)
add_callout(doc, "DECISION", "Evolve the existing Phase 2–6 architecture into one coherent, reviewable manual ChatGPT workflow. Reuse current package, validation, intelligence, comments, search, report, collection, and organization code. Do not replace the application with a greenfield design.", TEAL, LIGHT_TEAL)

add_table(doc, ["Document field", "Value"], [
    ("Application", "VideoHoarder desktop application"),
    ("Primary exchange", "Local package creation → manual ChatGPT upload → local result import"),
    ("Local AI", "Existing Ollama Ask Local AI remains available and separate"),
    ("Source specification", "Codex_Master_Spec_ChatGPT_Video_Intelligence_CHATGPT_ONLY.docx"),
    ("Implementation posture", "Incremental migration with backward compatibility"),
], [1.55, 4.9], 9)

doc.add_page_break()

doc.add_heading("1. Executive Decision", level=1)
doc.add_paragraph("The source specification contains valuable intelligence requirements, but it describes a more unified system than VideoHoarder currently has. VideoHoarder already implements much of the required foundation through several overlapping generations of package, import, validation, report, search, comment, collection, and local-AI code. The correct implementation is therefore consolidation and extension—not replacement.")
add_callout(doc, "CORE RULE", "No operation sends data directly to ChatGPT. VideoHoarder prepares transparent local files. The user reviews and manually uploads them, then manually downloads the returned JSON and imports it.", BLUE, LIGHT_BLUE)
doc.add_paragraph("The adapted target has four package modes sharing one envelope and one normalized result model: Per-Video Master Intelligence, Focused Topic/Question, Taxonomy Cleanup, and Same-Topic Cross-Video Synthesis. Existing legacy detailed-summary packages remain importable during migration.")

doc.add_heading("2. What VideoHoarder Already Has", level=1)
add_table(doc, ["Capability", "Current implementation", "Adaptation decision"], [
    ("Detailed transcript summaries", "build_chatgpt_summary_packages; legacy summary schema and canonical summary files", "Retain importer compatibility; move new exports to the unified envelope."),
    ("Package lifecycle", "Phase 2 folders, manifests, history, retry, archive, integrity and completeness audits", "Use as the canonical package-state engine."),
    ("Validation", "Package/video-ID checks, English checks, timestamp normalization, overlap and quality analysis", "Extend for unified schema, provenance, multipart results and preview."),
    ("Structured extraction", "Phase 4 content types/cards and structured intelligence", "Map normalized master fields into existing content-card outputs."),
    ("Search and reports", "Phase 5 search index, topic pages, related videos, dashboards and report regeneration", "Refresh only changed VIDEO_IDs after approved import."),
    ("Collections", "Full and focused collection packages with transcript sections and comments", "Refactor into Focused and Same-Topic package modes."),
    ("Comments", "Meaningfulness scoring, meaningful-comments transcript, report injection and ChatGPT excerpt", "Keep separate evidence provenance; add usefulness recommendation and counts."),
    ("Organization", "Taxonomy update, folder move preview, apply and undo", "Preserve review-first behavior; prohibit destructive AI authority."),
    ("Local retrieval/AI", "Phase 6 chunks, embeddings, evidence search and Ollama generation", "Keep as local-only companion; do not remove because the source document says ChatGPT-only."),
    ("Artifact tracking", "Artifact registry, manifests, library audits and portable metadata backup", "Extend to normalized intelligence and raw ChatGPT result lineage."),
], [1.35, 2.55, 2.55], 7.8)

doc.add_heading("3. Important Differences from the Source Specification", level=1)
add_bullet(doc, "VideoHoarder is an existing desktop application with mature workflows, not a new ChatGPT-only service.")
add_bullet(doc, "It has multiple overlapping package formats and import paths that must remain compatible while a unified format is introduced.")
add_bullet(doc, "The user performs the ChatGPT transfer manually; API keys, direct OpenAI calls, and automatic uploads are not part of this design.")
add_bullet(doc, "Ollama already powers a local question-answer feature. It remains local and separate instead of being removed.")
add_bullet(doc, "Folder organization is physical and potentially disruptive, so ChatGPT updates must continue through preview/apply/undo.")
add_bullet(doc, "The library can contain partial, cached, multilingual, or missing transcripts. Package behavior must report evidence coverage honestly.")

doc.add_heading("4. Revised End-to-End Flow", level=1)
code_block(doc, "SELECT VIDEOS / COLLECTION / TOPIC\n  -> LOCAL EVIDENCE ASSEMBLY\n  -> PREFLIGHT + PRIVACY REVIEW\n  -> PACKAGE CREATION (no upload)\n  -> USER MANUALLY UPLOADS TO CHATGPT\n  -> USER SAVES RETURNED JSON\n  -> INBOX DETECTION + RAW ARCHIVE\n  -> SCHEMA / ID / TIMESTAMP / EVIDENCE VALIDATION\n  -> NORMALIZATION + IMPORT PREVIEW / DIFF\n  -> USER APPROVES APPLY\n  -> ATOMIC PER-VIDEO COMMIT\n  -> SELECTIVE REPORT / SEARCH / KNOWLEDGE REFRESH\n  -> AUDIT, HISTORY, RETRY OR ROLLBACK")

doc.add_heading("4.1 Stage A — Selection", level=2)
add_bullet(doc, "Allow selection by unclassified status, newly downloaded videos, channel, category, explicit VIDEO_IDs, collection, Knowledge Center topic, search results, or test sample.")
add_bullet(doc, "Show exact candidate count, transcript availability, comments availability, estimated package size, and any locked/manual fields before export.")
add_bullet(doc, "Do not mark a video EXPORTED until the package and manifest are successfully written and validated locally.")

doc.add_heading("4.2 Stage B — Evidence Assembly", level=2)
add_bullet(doc, "Resolve identity from VIDEO_ID first. Never put fragile local paths into a ChatGPT package.")
add_bullet(doc, "Prefer timestamped English transcript; otherwise use timestamped source language, clean transcript, partial sections, then metadata-only evidence.")
add_bullet(doc, "Keep transcript evidence, official metadata, existing intelligence, and viewer comments in separate labeled blocks.")
add_bullet(doc, "Include prior accepted intelligence only when doing incremental enrichment, and label it existing_intelligence rather than transcript evidence.")

doc.add_heading("4.3 Stage C — Preflight and Privacy Review", level=2)
add_bullet(doc, "Display every field that will be included, counts and character totals by evidence source, and the exact output path.")
add_bullet(doc, "Warn when descriptions, transcripts, or comments appear to contain email addresses, phone numbers, credentials, or private names; allow exclude/redact/cancel.")
add_bullet(doc, "Confirm that local paths, cookies, API keys, logs, database files, media, and browser data are excluded.")
add_bullet(doc, "Generate a human-readable PACKAGE_README alongside JSON explaining manual upload and expected return filename.")

doc.add_heading("4.4 Stage D — Manual ChatGPT Exchange", level=2)
add_number(doc, "VideoHoarder writes the package, manifest, prompt/schema versions, and checksum locally.")
add_number(doc, "The user opens ChatGPT and uploads the reviewed package.")
add_number(doc, "ChatGPT returns JSON only, matching package_id and VIDEO_ID values exactly.")
add_number(doc, "The user downloads the result into the application root, results folder, or exchange inbox.")
add_number(doc, "VideoHoarder detects but does not apply the result until validation and preview succeed.")

doc.add_heading("4.5 Stage E — Validate, Preview, Apply", level=2)
add_bullet(doc, "Archive the raw returned file before normalization, including invalid returns.")
add_bullet(doc, "Validate JSON syntax, schema version, package ID, expected VIDEO_IDs, duplicates, unknown IDs, English-output rules, timestamps, evidence references, confidence ranges, and prohibited actions.")
add_bullet(doc, "Show field-level before/after differences and identify source, confidence, lock state, and downstream consequences such as folder moves.")
add_bullet(doc, "Apply each VIDEO_ID atomically. A failed video rolls back without preventing other valid videos from committing.")
add_bullet(doc, "Refresh only the changed video’s summary, content cards, report, search document, chunks, embeddings, related-video cache, topic pages, and collection views.")

doc.add_heading("5. Unified Package Envelope", level=1)
doc.add_paragraph("All new modes use one outer envelope so manifests, validation, history, retries, and imports do not need separate implementations.")
code_block(doc, '''{
  "package_format": "videohoarder_chatgpt_package",
  "package_version": "3.0",
  "schema_version": "3.0",
  "prompt_version": "vh-master-1",
  "package_id": "pkg_YYYYMMDD-HHMMSS_random",
  "mode": "MASTER|FOCUSED|TAXONOMY|SAME_TOPIC",
  "created_at": "ISO-8601",
  "selection": {"source": "...", "video_ids": []},
  "privacy_review": {"completed": true, "excluded_fields": []},
  "instructions": {},
  "videos": [],
  "expected_output": {},
  "checksums": {}
}''')
add_bullet(doc, "Use a collision-resistant package ID; the current timestamp-only approach can collide when packages are created in the same second.")
add_bullet(doc, "Manifest records expected IDs, source hashes, evidence sizes, transcript/comment status, part numbers, and required output filenames.")
add_bullet(doc, "Package generation is idempotent for the same selection and source hashes unless the user explicitly creates a new revision.")

doc.add_heading("6. Exact Information Manually Sent to ChatGPT", level=1)
add_table(doc, ["Evidence group", "Fields", "Rules"], [
    ("Identity", "VIDEO_ID, original title, current approved title, video URL", "VIDEO_ID is mandatory and immutable. No local file path."),
    ("Official metadata", "Channel, channel URL when known, description, upload date, download date, duration, official category", "Description is labeled metadata, not transcript evidence."),
    ("Current organization", "Category, subcategory, tags, primary topic, manual-lock indicators", "Used as prior state; ChatGPT may propose but not silently override locked fields."),
    ("Transcript", "Status, language, timestamped transcript, clean transcript or focused sections", "No silent truncation. Split/continue when full evidence is required."),
    ("Existing sections", "Chapter/section titles, start/end timestamps, summaries", "Label as prior structured intelligence when not directly sourced."),
    ("Viewer comments", "Selected meaningful comments, comment IDs/hashes, vote/author fields when retained, aggregate count", "Audience-generated and unverified; separate from transcript."),
    ("Existing intelligence", "Prior accepted summary/taxonomy/cards when incremental analysis is selected", "Optional and explicitly labeled."),
    ("Task controls", "Mode, topic/question, required schema, anti-hallucination rules", "No external facts unless a separately approved verification mode exists."),
], [1.2, 3.0, 2.25], 7.7)
add_callout(doc, "NOT SENT", "Video/audio files, thumbnails, Windows paths, SQLite database, cookies, credentials, API keys, browser data, logs, unrelated library records, Ollama settings, and other local files.", RED, "FEECEB")

doc.add_heading("7. Package Modes", level=1)
doc.add_heading("7.1 MASTER — Per-Video Final Intelligence", level=2)
add_bullet(doc, "Replaces new use of the legacy detailed-summary exporter while retaining its importer compatibility.")
add_bullet(doc, "Includes complete evidence where practical; large inputs are split into per-video parts and followed by a merge package.")
add_bullet(doc, "Returns full title, taxonomy, summaries, timeline, named items, cards, claims, warnings, sources, glossary, Q&A, comments intelligence, evidence provenance, and quality fields.")

doc.add_heading("7.2 FOCUSED — Topic or Question", level=2)
add_bullet(doc, "Uses Phase 6 retrieval to select relevant sections plus surrounding context and any repeated occurrences from the same video.")
add_bullet(doc, "States what evidence was omitted and never presents a focused package as complete video coverage.")
add_bullet(doc, "Returns topic-specific findings and timestamp ranges without overwriting unrelated master intelligence.")

doc.add_heading("7.3 TAXONOMY — Classification and Tag Cleanup", level=2)
add_bullet(doc, "Combines current taxonomy package and JSONL tag-cleanup workflows behind one UI action and shared schema.")
add_bullet(doc, "Uses metadata plus limited evidence needed for classification; it does not pretend to be a detailed summary.")
add_bullet(doc, "Returns title proposal, category, subcategory, primary topic, canonical tags, concepts, search phrases, synonyms, questions answered, not_about, basis, and confidence.")

doc.add_heading("7.4 SAME_TOPIC — Cross-Video Synthesis and Clip Plan", level=2)
add_bullet(doc, "Starts from a collection, topic page, search, selected videos, or detected topic.")
add_bullet(doc, "Collapses exact/near-duplicate candidates by default and shows why each candidate matched.")
add_bullet(doc, "Returns per-video topic summaries, all relevant ranges, unique contributions, agreements, differences, contradictions, synthesis, and a reviewable VIDEO_ID-based clip plan.")
add_bullet(doc, "Never cuts or merges automatically. Validated plans are previewed before being passed to the existing Clip Studio resolver.")

doc.add_heading("8. Unified Normalized Intelligence", level=1)
doc.add_paragraph("Do not force every old importer and report to understand every ChatGPT variant. Normalize old and new returns into one internal model, then project that model into the existing summary, content-card, search, report, and taxonomy structures.")
add_table(doc, ["Domain", "Normalized fields"], [
    ("Identity/title", "video_id, original_title, approved_english_title, short_title, source_language, title_translated"),
    ("Taxonomy/search", "category, subcategory, primary_topic, secondary_topics, canonical_tags, concepts, search_phrases, synonyms, questions_answered, not_about, content_types"),
    ("Summary", "executive_summary, detailed_summary, key_takeaways, important_facts, important_numbers"),
    ("Timeline", "timeline_sections, topic_occurrences, compact_chapters, timestamp_coverage"),
    ("Structured content", "how_to_guides, remedies, recipes, comparisons, rankings, recommendations, products_services_tools"),
    ("Evidence-sensitive", "claims, warnings_safety, sources_references, glossary, qa, entities, things_discussed"),
    ("Comments", "viewer_comment_intelligence with separate evidence IDs and cautions"),
    ("Quality/provenance", "evidence_provenance, analysis_quality, field_confidence, transcript_status, classification_basis"),
    ("Optional commentary", "chatgpt_perspective_ps, always labeled as not source-extracted fact"),
], [1.55, 4.9], 8.4)

doc.add_heading("8.1 Evidence Object", level=2)
code_block(doc, '''{
  "source_type": "timestamped_transcript|transcript|description|metadata|chapter|viewer_comment|existing_intelligence",
  "source_id": "stable local evidence ID",
  "start_timestamp": "HH:MM:SS",
  "end_timestamp": "HH:MM:SS",
  "source_excerpt": "short supporting excerpt",
  "confidence": 0.0
}''')
add_bullet(doc, "Important facts, claims, warnings, recommendations, steps, quantities, and named items must carry evidence when available.")
add_bullet(doc, "Evidence priority: timestamped transcript → transcript → official metadata/description → accepted existing intelligence → viewer comment.")

doc.add_heading("9. Validation and Import Rules", level=1)
add_table(doc, ["Validation gate", "Required behavior"], [
    ("File and schema", "Valid JSON; known format/schema; prompt version recorded; raw return retained."),
    ("Package identity", "package_id matches manifest; part/merge relationships are complete."),
    ("Video identity", "Exact real VIDEO_ID; reject placeholders, duplicates and unknown IDs."),
    ("Coverage", "Expected IDs compared with returned IDs; partial results may be imported only when explicitly allowed."),
    ("Language", "Required display fields are English; proper nouns and source-language aliases remain allowed."),
    ("Timestamps", "Parse and normalize; start ≤ end; within duration tolerance; no invented ranges; repeated topic ranges allowed."),
    ("Evidence", "Evidence source exists in the exported package; excerpts/ranges are plausible and source types are valid."),
    ("Confidence", "All confidence values are numeric 0–1; metadata-only analysis is prevented from claiming full transcript coverage."),
    ("Locks/authority", "Locked manual fields and destructive instructions cannot be applied without explicit user action."),
    ("Safety", "Speaker claims, viewer experiences and optional ChatGPT perspective remain separately labeled."),
], [1.35, 5.1], 8.2)

doc.add_heading("9.1 Import Preview", level=2)
add_bullet(doc, "Show current value, proposed value, source type, confidence, lock state, and selected apply action per field.")
add_bullet(doc, "Provide Accept All Safe, Accept Selected, Reject Video, Create Retry Package, and Save Without Applying.")
add_bullet(doc, "Show report/search rebuilds and any proposed folder movement separately.")
add_bullet(doc, "Folder moves remain a second preview/apply action after metadata import.")

doc.add_heading("9.2 Atomicity and Undo", level=2)
add_bullet(doc, "Write STARTED, COMMITTED, ROLLED_BACK, or FAILED import transaction state per VIDEO_ID.")
add_bullet(doc, "Back up the database row and affected canonical intelligence/report files before commit.")
add_bullet(doc, "On failure, restore only that VIDEO_ID. Never leave a half-updated report, search record, card set, or folder path.")
add_bullet(doc, "Undo restores pre-import metadata and intelligence; it does not delete media.")

doc.add_heading("10. Storage and Lineage", level=1)
code_block(doc, '''data/chatgpt/
  packages/                 # outgoing packages + README + manifest
  results/raw/              # immutable user-returned files
  results/normalized/       # normalized per-package results
  retry/                    # generated retry packages
  archive/                  # completed package bundles
  rejected/                 # invalid files + validation report
  previews/                 # field diffs awaiting approval
  transactions/             # per-VIDEO_ID import state and rollback data
  package_history.csv

per-video _data/
  <VIDEO_ID>.chatgpt_summary.json
  <VIDEO_ID>.chatgpt_summary.txt
  <VIDEO_ID>.intelligence.json
  <VIDEO_ID>.intelligence_provenance.json
  <VIDEO_ID>.comments_intelligence.json''')
add_bullet(doc, "Keep the existing canonical summary filenames and add schema_version, source_hashes, package_id, prompt_version, imported_at, and normalized_hash.")
add_bullet(doc, "Cross-video synthesis lives separately and must not overwrite per-video intelligence.")
add_bullet(doc, "Artifact Registry records raw result, normalized result, report, search record, chunk/embedding versions, comments intelligence and stale status.")

doc.add_heading("11. Downstream Processing After Approved Import", level=1)
add_table(doc, ["Stage", "Action"], [
    ("Store", "Save normalized intelligence and the canonical summary for the VIDEO_ID."),
    ("Project", "Map structured content into existing Phase 4 content-card outputs."),
    ("Report", "Regenerate the individual HTML report with source labels and clickable local timestamps."),
    ("Search", "Update the Phase 5 search document, taxonomy vocabulary, topic pages and related-video data."),
    ("Retrieval", "Rebuild Phase 6 chunks and embeddings only when relevant source/intelligence hashes changed."),
    ("Collections", "Update collection and Knowledge Center views that reference the VIDEO_ID."),
    ("Organize", "Prepare—but do not automatically apply—the category/channel folder move preview."),
    ("Complete", "Archive only when all required outputs pass completeness checks."),
], [1.0, 5.45], 8.2)

doc.add_heading("12. UI Changes", level=1)
doc.add_paragraph("Keep the existing ChatGPT Processing page and its maintenance buttons, but reorganize it into a guided workspace.")
add_table(doc, ["Panel", "Controls and information"], [
    ("1. Select", "Mode, selection source, filters, candidate list, include/exclude, lock indicators."),
    ("2. Review evidence", "Per-video transcript/comments/metadata coverage, sizes, privacy warnings, omissions."),
    ("3. Create package", "Token/character estimate, split plan, output folder, Create and Open Folder."),
    ("4. Manual exchange", "Upload instructions, expected return name, Open ChatGPT, Open Inbox—without automatic upload."),
    ("5. Import", "Detected results, validation status, raw/normalized files, preview/diff, retry."),
    ("6. Apply", "Selected field updates, report/search refresh, move preview, undo."),
    ("Maintenance", "Legacy Migration—Sync ChatGPT Folders; Phase 2 Integrity & Completeness Audit; Rebuild Comment Intelligence & Comments Transcript."),
], [1.4, 5.05], 8.5)
add_callout(doc, "STATUS LANGUAGE", "Use precise states: Not Created, Ready for Manual Upload, Awaiting Returned File, Validation Failed, Ready for Review, Partially Imported, Applied, Retry Required, Archived. Never say “Sent to ChatGPT” because the app does not perform the upload.", AMBER, "FEF0C7")

doc.add_heading("12.1 Existing ChatGPT Title Rename Workflow", level=2)
doc.add_paragraph("Physical file renaming from imported ChatGPT titles already exists and must be preserved. It is not a new feature. The unified review/apply workspace will invoke the existing safe rename workflow after title approval.")
add_table(doc, ["Existing behavior", "Decision"], [
    ("Import meaningful_title / approved_english_title", "Keep; normalize both fields into the approved title."),
    ("Update SQLite clean_title", "Keep; apply only after validation and field-lock review."),
    ("Rename physical media file", "Keep existing safe rename and conflict checks."),
    ("Rename report and related marker/support artifacts", "Keep and re-detect files after rename."),
    ("Rename/move the video folder", "Keep existing path update and report-rebuild behavior."),
    ("Rebuild HTML report and playable media link", "Keep as a required post-rename completeness check."),
], [2.55, 3.9], 8.3)
add_callout(doc, "IMPLEMENTATION", "Reuse apply_chatgpt_titles_and_rename(), the one-video ChatGPT rename workflow, safe rename helpers, path updates and rebuild_report_after_rename(). Add only a clearer preview, manual-lock check and transaction/undo boundary around them.", TEAL, LIGHT_TEAL)

doc.add_heading("13. Comments, Claims and Safety", level=1)
add_bullet(doc, "Comment intelligence remains a separate evidence layer. Rebuilding it is local processing and does not itself send data.")
add_bullet(doc, "Before downloading/packaging comments, calculate a usefulness recommendation with reasons; preserve manual override and existing comments.")
add_bullet(doc, "Health, finance, legal, political, scientific, and product claims retain speaker attribution, timestamp, evidence and verification state.")
add_bullet(doc, "Never invent dosage, contraindications, warnings, prices, versions, quantities, sources, list items, steps, or timestamps.")
add_bullet(doc, "Optional ChatGPT perspective is stored and displayed in a visually separate labeled section and never outranks transcript evidence.")

doc.add_heading("14. Related Videos, Duplicates and Same-Topic Intelligence", level=1)
add_bullet(doc, "Reuse Phase 5 related-video scoring and Phase 6 semantic retrieval as candidate generators.")
add_bullet(doc, "Add stable duplicate-group records using VIDEO_ID, media fingerprints/hashes where available, duration, title/channel, transcript similarity and intelligence similarity.")
add_bullet(doc, "Duplicate detection recommends a canonical source but never deletes or archives variants automatically.")
add_bullet(doc, "Related-video ranking collapses duplicate groups and rewards complementary/unique information, not only similarity.")
add_bullet(doc, "Same-topic synthesis keeps agreements, contradictions and unique contributions separate from per-video facts.")

doc.add_heading("14.1 Exact Status of Manual Operations", level=2)
add_table(doc, ["Operation", "Current status", "Final treatment"], [
    ("Manual deletion of duplicate videos", "Partially available through duplicate audits, Mark Delete and existing deletion/cleanup controls.", "Add a dedicated duplicate-group review screen; deletion remains explicit and confirmed."),
    ("Choose a canonical duplicate", "Not yet a complete integrated workflow.", "Add side-by-side comparison, recommendation and manual canonical selection. Selection never auto-deletes variants."),
    ("Delete/archive non-canonical variants", "General manual controls exist; not integrated with canonical selection.", "Add Keep, Archive, Mark Delete and Delete actions with confirmation and audit history."),
    ("Cut/merge from ChatGPT clip plan", "Clip Studio and local VIDEO_ID-based merging already exist; ChatGPT plan handoff does not.", "Add merge_clip_plan validation, preview, inclusion/reordering and Send Merge Plan to Clip Studio."),
    ("Move folders from ChatGPT output", "Already implemented through taxonomy move preview, apply, history and undo.", "Preserve. ChatGPT proposes category; VideoHoarder calculates the safe local destination."),
], [1.65, 2.35, 2.45], 7.5)
add_callout(doc, "NO DESTRUCTIVE AI AUTHORITY", "ChatGPT may recommend a title, category, canonical copy or clip plan. Only the user can approve physical renaming, organization moves, cutting/merging, archiving or deletion.", RED, "FEECEB")

doc.add_heading("15. Local Ollama Boundary", level=1)
doc.add_paragraph("The source document says ChatGPT is the only AI engine, but that is not appropriate as a destructive migration rule for VideoHoarder because the app already has a working local Ollama evidence-answer feature. The adapted rule is:")
add_callout(doc, "BOUNDARY", "ChatGPT is the only engine used for the manual package/import intelligence workflow defined here. Existing Ollama remains available for local Ask Local AI and embeddings/retrieval support. It must not silently write authoritative ChatGPT intelligence fields or bypass import review.", TEAL, LIGHT_TEAL)

doc.add_heading("16. Implementation Plan", level=1)
add_table(doc, ["Phase", "Work", "Primary outcome"], [
    ("0 — Baseline", "Inventory package/import/report/search/comment paths; freeze fixtures; back up DB/config.", "Documented reuse map and regression safety."),
    ("1 — Schema", "Add v3 envelope, normalized model, schema/prompt registry and adapters for legacy returns.", "One internal interpretation for old and new results."),
    ("2 — Export", "Create shared evidence builder, privacy preflight, estimates, splitting/merge and four package modes.", "Transparent manual-upload packages."),
    ("3 — Import", "Raw archive, validation, normalization, field diff, locks, atomic per-video transactions and retry.", "Safe reviewable imports."),
    ("4 — Projection", "Map normalized data to summaries, Phase 4 cards, reports, Phase 5 search and Phase 6 chunks/embeddings.", "No parallel intelligence silo."),
    ("5 — UI", "Guided ChatGPT workspace, truthful states, package dashboard, preview/apply/undo, maintenance grouping.", "Understandable end-to-end operation."),
    ("6 — Advanced", "Same-topic synthesis, clip plans, duplicate groups, related diversity, comment usefulness, portable export.", "Cross-video intelligence and safe Clip Studio handoff."),
    ("7 — Hardening", "Locks, numbered DB migrations, orphan audit, performance tests, crash/recovery and package safety.", "Production reliability at 10,000+ videos."),
], [0.8, 3.7, 1.95], 7.6)

doc.add_heading("16.1 Concrete Code Refactoring", level=2)
add_bullet(doc, "Extract common evidence assembly from build_chatgpt_summary_packages, phase6_taxonomy_video_payload and phase6_collection_package.")
add_bullet(doc, "Use Phase 2 package folders, history, manifests, retry and archive as shared lifecycle services.")
add_bullet(doc, "Extend phase2_normalize_result_payload and normalize_chatgpt_result_video into versioned adapters feeding the v3 model.")
add_bullet(doc, "Reuse validate_and_normalize_sections plus Phase 3 language/timestamp/quality checks inside the new validator.")
add_bullet(doc, "Route normalized structured content through Phase 4 extraction/writers rather than duplicating cards.")
add_bullet(doc, "Use Phase 5 incremental hashes and Phase 6 source hashes to refresh only stale downstream artifacts.")
add_bullet(doc, "Retain phase6_preview_category_moves, phase6_apply_previewed_moves and phase6_undo_last_organization as the physical-move authority.")

doc.add_heading("17. Database and Migration Changes", level=1)
add_bullet(doc, "Introduce numbered, restart-safe migrations and a schema_migrations history table before adding new columns/tables.")
add_bullet(doc, "Add package revisions, prompt/schema versions, raw/normalized result hashes, per-video import transaction state, field provenance, manual locks, duplicate groups, and stale flags.")
add_bullet(doc, "Back up the database before each migration and validate row counts, VIDEO_ID uniqueness, foreign references, and artifact paths afterward.")
add_bullet(doc, "Do not scatter new ALTER TABLE statements across runtime functions.")

doc.add_heading("18. Compatibility and Non-Destructive Rules", level=1)
add_bullet(doc, "Continue importing old detailed-summary, taxonomy-only and JSONL tag-cleanup results.")
add_bullet(doc, "Never overwrite or delete the user’s original package/result files during migration; copy into canonical storage and record source path/hash.")
add_bullet(doc, "ChatGPT output can propose metadata, taxonomy, organization and canonical-source changes, but cannot delete media or bypass confirmation.")
add_bullet(doc, "Manual field locks win over ChatGPT output. An explicit unlock/override is required.")
add_bullet(doc, "Legacy Migration—Sync ChatGPT Folders stays available until all legacy packages are mapped and audited.")

doc.add_heading("19. Testing and Acceptance", level=1)
add_table(doc, ["Test group", "Minimum cases"], [
    ("Legacy compatibility", "Old summary result, old taxonomy result, JSONL cleanup, missing package ID, partial return."),
    ("Identity", "Valid ID, placeholder ID, duplicate ID, unknown ID, package mismatch."),
    ("Transcript/language", "English, Hindi/multilingual, missing, partial, unusable, oversized multipart."),
    ("Structured content", "Remedy/health, recipe, technical tutorial, comparison/ranking, comments, repeated topic."),
    ("Validation", "Bad JSON, out-of-range timestamp, fabricated evidence ID, invalid confidence, prohibited delete request."),
    ("Transactions", "Injected failure mid-video, batch with one bad video, rollback, repeated idempotent import."),
    ("Concurrency", "Download/import/report rebuild/move compete for the same VIDEO_ID; stale-lock recovery."),
    ("Scale", "100, 1,000 and 10,000+ synthetic records; incremental work avoids transcript all-pairs on page load."),
    ("Privacy/package", "No local paths, keys, cookies, DB, logs or media; warning/redaction controls work."),
    ("Same-topic/Clip Studio", "Multiple ranges per video, duplicate collapse, NOT_MERGE_READY when timestamps are missing."),
], [1.35, 5.1], 8.1)

doc.add_heading("20. Definition of Done", level=1)
add_bullet(doc, "A user can see exactly what will be manually uploaded before a file is created.")
add_bullet(doc, "All four package modes use the same lifecycle, versioning, manifest, raw archive, normalization, preview and audit infrastructure.")
add_bullet(doc, "Old result formats still import through tested adapters.")
add_bullet(doc, "No ChatGPT result changes a locked field, deletes media, or moves folders without explicit review.")
add_bullet(doc, "Every important normalized field can be traced to package, prompt/schema version, source type and evidence where available.")
add_bullet(doc, "Per-video imports are atomic, retryable, reversible and selectively refresh downstream artifacts.")
add_bullet(doc, "The local Ollama assistant remains local and cannot bypass the reviewed ChatGPT intelligence pipeline.")
add_bullet(doc, "Integrity, orphan, migration, package-completeness and performance tests pass on a copy of the real library before release.")

doc.add_heading("Appendix A — Feature Disposition", level=1)
add_table(doc, ["Source-spec feature", "Disposition for VideoHoarder"], [
    ("Unified master schema", "Implement through normalized v3 model; keep adapters for old schemas."),
    ("Full/focused/taxonomy modes", "Implement as shared envelope; map collection full/focused and taxonomy exporters."),
    ("Direct ChatGPT analysis", "Manual file exchange only; no API upload."),
    ("Knowledge Center", "Reuse and enrich existing Phase 5/6 outputs."),
    ("Ask ChatGPT about selected videos", "Generate a Focused package for manual ChatGPT; keep Ask Local AI via Ollama separately."),
    ("Find Timestamp", "Reuse Phase 5/6 retrieval; add repeated occurrences and provenance."),
    ("Related videos", "Extend existing implementation with duplicate collapse and novelty/diversity."),
    ("Duplicate detection", "Extend existing audits into stable groups and reviewable canonical recommendations."),
    ("Same-topic/clip plan", "Add as separate cross-video layer; validate before Clip Studio preview."),
    ("External verification", "Defer to a separate opt-in workflow; never mix with source-only extraction."),
    ("Local alternative AI", "Preserve existing Ollama; outside authoritative ChatGPT import path."),
    ("ChatGPT-based physical renaming", "Already implemented; preserve safe media/report/folder rename and post-rename report rebuild."),
    ("ChatGPT taxonomy folder movement", "Already implemented with preview, apply, history and undo; retain local destination calculation."),
    ("Canonical duplicate selection", "Add as an advanced reviewed workflow; recommendation only, never automatic deletion."),
    ("ChatGPT clip-plan handoff", "Add validation and preview on top of the existing Clip Studio merge engine."),
], [2.25, 4.2], 8)

doc.add_page_break()
doc.add_heading("Appendix B — Operational Meaning of Existing Maintenance Buttons", level=1)
add_bullet(doc, "Legacy Migration — Sync ChatGPT Folders: copy/map older packages and results into canonical Phase 2 storage; create missing manifests; no upload.")
add_bullet(doc, "Phase 2 Package Integrity & Completeness Audit: read-only verification of manifests, expected/returned VIDEO_IDs, missing results, duplicates and archive readiness; no upload.")
add_bullet(doc, "Rebuild Comment Intelligence & Comments Transcript: locally re-score downloaded comments and rebuild meaningful-comment artifacts; only later package creation can include selected comment evidence.")

# Keep headings with content and normalize all body runs.
for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        set_repeat_keep(p, keep_next=True, keep_lines=True)
    for run in p.runs:
        if not run.font.name:
            set_font(run)

doc.core_properties.title = "VideoHoarder ChatGPT Video Intelligence Adapted Master Specification"
doc.core_properties.subject = "Application-specific implementation design for manual ChatGPT package processing"
doc.core_properties.author = "VideoHoarder Project"
doc.core_properties.keywords = "VideoHoarder, ChatGPT, video intelligence, manual package, implementation specification"
doc.save(OUT)
print(OUT)
