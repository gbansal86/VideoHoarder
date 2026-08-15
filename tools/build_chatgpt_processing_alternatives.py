from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Users\gaurav.bansal\Documents\ChatGPT\VideoHoarder")
OUT = ROOT / "ChatGPT_Processing_UI_Design_Alternatives.docx"
BLUE = RGBColor(45, 115, 195)
NAVY = RGBColor(19, 48, 82)
GRAY = RGBColor(97, 112, 130)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)
styles["Normal"].paragraph_format.space_after = Pt(5)
for name, size, color in [("Heading 1", 16, NAVY), ("Heading 2", 12, BLUE)]:
    style = styles[name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(10)
    style.paragraph_format.space_after = Pt(5)
    style.paragraph_format.keep_with_next = True

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ChatGPT Processing UI Design Alternatives")
r.font.name = "Aptos Display"; r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = NAVY
subtitle = doc.add_paragraph("VideoHoarder - Designs retained for future reference")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in subtitle.runs:
    r.font.name = "Aptos"; r.font.size = Pt(11); r.font.color.rgb = GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Selected implementation: Design 1 - Sidebar Command Center")
r.font.bold = True; r.font.color.rgb = BLUE
doc.add_paragraph("This document intentionally preserves the four non-selected layouts. They are not part of the current implementation specification, but can be reconsidered later if package volume, user roles, or workflow needs change.")

concepts = [
    (2, "Guided Top-Step Console", "Best for users who want the application to lead them through one package at a time.", "The six-stage flow stays prominent, with fewer competing panels. It is strong for first-time and occasional use.", "History, coverage, and advanced tools are less continuously visible and require tab changes."),
    (3, "Three-Pane Package Workbench", "Best for high-volume package creation and feature-enrichment campaigns.", "Package list, current work, and inspector remain visible together; excellent for evidence and coverage decisions.", "More complex and best suited to wider desktop windows; can feel dense for simple one-off tasks."),
    (4, "Package Lifecycle Board", "Best for monitoring many packages at different stages.", "Makes Prepare, Awaiting ChatGPT, Validate/Review, and Applied/Archived states instantly understandable.", "Less suitable for detailed evidence review and field-level editing; another detail view is still required."),
    (5, "Dense Operations and Coverage Console", "Best for expert users with thousands of videos and frequent incremental processing.", "Combines search, feature coverage, package state, and history in a compact data-oriented workspace.", "Highest learning curve and least approachable for users who prefer a guided workflow."),
]

for position, (idx, title, best, strength, tradeoff) in enumerate(concepts):
    doc.add_page_break()
    doc.add_heading(f"Design {idx} - {title}", level=1)
    image = ROOT / "tools" / f"chatgpt-processing-concept-{idx}.png"
    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(image), width=Inches(6.65))
    caption = doc.add_paragraph(f"Figure {position + 1} - Design {idx}: {title}")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.size = Pt(8); run.font.color.rgb = GRAY
    for label, value in [("Best for", best), ("Strength", strength), ("Tradeoff", tradeoff)]:
        p = doc.add_paragraph()
        r = p.add_run(label + ": "); r.font.bold = True; r.font.color.rgb = BLUE
        p.add_run(value)

doc.add_page_break()
doc.add_heading("When to Revisit These Alternatives", level=1)
doc.add_paragraph("Revisit Design 2 if the primary need becomes guided, low-frequency use. Revisit Design 3 for a specialist workflow with frequent package construction and evidence review. Revisit Design 4 if package monitoring becomes the main daily task. Revisit Design 5 for expert teams operating very large libraries.")

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.add_run("VIDEOHOARDER | CHATGPT PROCESSING DESIGN ALTERNATIVES").font.size = Pt(8)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Design alternatives - not selected for current implementation").font.size = Pt(8)

doc.core_properties.title = "ChatGPT Processing UI Design Alternatives"
doc.core_properties.subject = "VideoHoarder alternative UI concepts"
doc.core_properties.author = "VideoHoarder Project"
doc.save(OUT)
print(OUT)
