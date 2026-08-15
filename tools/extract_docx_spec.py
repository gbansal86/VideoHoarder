from docx import Document

path = r"C:\Users\gaurav.bansal\Downloads\Codex_Master_Spec_ChatGPT_Video_Intelligence_CHATGPT_ONLY.docx"
doc = Document(path)
print(f"PARAGRAPHS={len(doc.paragraphs)} TABLES={len(doc.tables)} SECTIONS={len(doc.sections)}")
for index, paragraph in enumerate(doc.paragraphs):
    if paragraph.text.strip():
        print(f"P{index} [{paragraph.style.name}]: {paragraph.text}")
for table_index, table in enumerate(doc.tables):
    print(f"TABLE {table_index}")
    for row in table.rows:
        print(" | ".join(cell.text.replace("\n", " / ") for cell in row.cells))
